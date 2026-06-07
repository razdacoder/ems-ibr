"""Shared institution-branding helpers for generated documents.

Both the attendance-sheet generator (``ems.views``) and the seat-arrangement
exporter (``ems.csv_gen``) print the same institution header on every Word
document: the configured logo, institution name/heading/address/contact, and
the active session & semester. This module centralises that so the two stay in
sync and read from ``SystemSettings`` rather than a hard-coded static path.
"""

import os

from django.conf import settings as django_settings
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def resolve_logo_path(settings_obj) -> str | None:
    """Filesystem path to the branding logo, or None if none is usable.

    Prefers the uploaded ``SystemSettings.logo``; falls back to the bundled
    ``static/assets/images/logo.png`` so existing deployments keep their mark.
    """
    logo = getattr(settings_obj, "logo", None)
    if logo:
        try:
            if logo.storage.exists(logo.name):
                return logo.path
        except (ValueError, NotImplementedError):
            pass
    fallback = os.path.join(
        django_settings.BASE_DIR, "static", "assets", "images", "logo.png"
    )
    return fallback if os.path.exists(fallback) else None


def add_document_branding(doc, settings_obj) -> None:
    """Prepend the institution header (logo + metadata + session) to ``doc``.

    Mirrors the previous inline header layout: left-aligned logo, then a
    centred institution block, then the session/semester line.
    """
    # Logo (or a text placeholder when no image is available).
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_path = resolve_logo_path(settings_obj)
    if logo_path:
        try:
            run = logo_paragraph.add_run()
            run.add_picture(logo_path, width=Inches(1.0))
        except Exception:
            placeholder = logo_paragraph.add_run("[SCHOOL LOGO]")
            placeholder.bold = True
    else:
        placeholder = logo_paragraph.add_run("[SCHOOL LOGO]")
        placeholder.bold = True

    # Institution block — only render the lines that are configured.
    exam_heading = (getattr(settings_obj, "exam_heading", "") or "").strip()
    name = (getattr(settings_obj, "institution_name", "") or "").strip()
    address = (getattr(settings_obj, "institution_address", "") or "").strip()
    email = (getattr(settings_obj, "contact_email", "") or "").strip()
    phone = (getattr(settings_obj, "contact_phone", "") or "").strip()

    if exam_heading:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(exam_heading.upper())
        r.bold = True
        r.font.size = Pt(12)

    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name.upper())
        r.bold = True
        r.font.size = Pt(14)

    contact_bits = [bit for bit in (address, email, phone) if bit]
    if contact_bits:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, bit in enumerate(contact_bits):
            run = p.add_run(bit)
            if i < len(contact_bits) - 1:
                run.add_break()

    # Session / semester line.
    session_paragraph = doc.add_paragraph()
    session_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    session_run = session_paragraph.add_run(
        f"SESSION: {settings_obj.session}"
    )
    session_run.bold = True
    session_run.add_break()
    semester_run = session_paragraph.add_run(
        f"SEMESTER: {settings_obj.semester}"
    )
    semester_run.bold = True

    # Spacer before the document body.
    doc.add_paragraph()
