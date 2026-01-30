import csv
from datetime import datetime
import io
import zipfile
import os
from django.http import HttpRequest, HttpResponse
from django.conf import settings

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from .models import Distribution, TimeTable, SeatArrangement, SystemSettings


def export_department_timetable(request: HttpRequest) -> HttpResponse:
    department = request.user.department
    filename = f'{department.name}-Timetable.csv'
    timetables = TimeTable.objects.filter(class_obj__department=department)
    response = HttpResponse(
        content_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename={filename}'},
    )
    writer = csv.writer(response)
    writer.writerow(
        ["Date", "Level", "Exam Type" "Course Code", "Course Title", "Period"])
    for timetable in timetables:
        writer.writerow([timetable.date.strftime('%A %d, %B %Y'), timetable.class_obj.name,
                        timetable.course.exam_type, timetable.course.code, timetable.course.name, timetable.period])

    return response


def export_distribution(request: HttpRequest) -> HttpResponse:
    date = request.GET.get("date")
    period = request.GET.get("period")

    if date is None:
        date = Distribution.objects.values_list(
            "date", flat=True).distinct().order_by("date").first()

    if period is None:
        period = "AM"

    # Handle case where no date is available
    if date is None:
        # Return empty CSV with error message
        response = HttpResponse(
            content_type="text/csv",
            headers={"Content-Disposition": 'attachment; filename="no-distribution-data.csv"'},
        )
        writer = csv.writer(response)
        writer.writerow(["Error: No distribution data available. Please generate distribution first."])
        return response

    if isinstance(date, str):
        date_obj = datetime.strptime(date, "%Y-%m-%d").date()
    else:
        date_obj = date

    distributions = Distribution.objects.filter(date=date_obj, period=period)

    filename = f"{date_obj.strftime('%A %d, %B %Y')} - {period}-Distribution.csv"

    response = HttpResponse(
        content_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
    writer = csv.writer(response)
    writer.writerow(["Hall", "Class Name", "No. Of Students"])
    for distribution in distributions:
        for item in distribution.items.all():
            cls = item.schedule.class_obj
            writer.writerow(
                [distribution.hall.name, f"{cls.department.slug} {cls.name}", item.no_of_students])
    return response


def export_arrangements(request: HttpRequest) -> HttpResponse:
    date = request.GET.get("date")
    period = request.GET.get("period")

    if date is None:
        date = SeatArrangement.objects.values_list(
            "date", flat=True).distinct().order_by("date").first()
    if period is None:
        period = "AM"

    # Handle case where no date is available
    if date is None:
        response = HttpResponse(
            content_type="application/zip",
            headers={"Content-Disposition": 'attachment; filename="no-arrangement-data.zip"'},
        )
        in_memory_zip = io.BytesIO()
        with zipfile.ZipFile(in_memory_zip, 'w') as zip_file:
            zip_file.writestr("error.txt", "Error: No arrangement data available. Please generate seat arrangements first.")
        response.write(in_memory_zip.getvalue())
        return response

    # Normalize date
    if isinstance(date, str):
        date_obj = datetime.strptime(date, "%Y-%m-%d").date()
    else:
        date_obj = date

    # Get system settings for session and semester
    settings_obj = SystemSettings.objects.first()
    if not settings_obj:
        settings_obj = SystemSettings.objects.create(
            session='2024/2025', semester='1st Semester')

    # Fetch arrangements across all halls for the date/period; include only placed students
    arrangements_qs = (
        SeatArrangement.objects
        .filter(date=date_obj, period=period, seat_number__isnull=False)
        .select_related('student', 'course', 'cls', 'hall')
        .order_by('course__name', 'hall__name', 'student__matric_no')
    )

    if not arrangements_qs.exists():
        response = HttpResponse(
            content_type="application/zip",
            headers={"Content-Disposition": 'attachment; filename="no-arrangement-data.zip"'},
        )
        in_memory_zip = io.BytesIO()
        with zipfile.ZipFile(in_memory_zip, 'w') as zip_file:
            zip_file.writestr("error.txt", "Error: No placed seat arrangements found for the specified criteria.")
        response.write(in_memory_zip.getvalue())
        return response

    # Group by course, then by hall
    grouped = {}
    for arr in arrangements_qs:
        course_id = arr.course_id
        hall_id = arr.hall_id
        if course_id not in grouped:
            grouped[course_id] = {
                'course': arr.course,
                'halls': {}
            }
        if hall_id not in grouped[course_id]['halls']:
            grouped[course_id]['halls'][hall_id] = {
                'hall': arr.hall,
                'cls': arr.cls,
                'students': []
            }
        grouped[course_id]['halls'][hall_id]['students'].append(arr)

    def safe_name(name: str) -> str:
        return ''.join(c if c.isalnum() or c in ' -_.' else '_' for c in name)

    # Create a zip containing DOCX attendance sheets grouped by course folders
    in_memory_zip = io.BytesIO()
    with zipfile.ZipFile(in_memory_zip, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for course_id, course_data in grouped.items():
            course = course_data['course']
            course_folder = f"{safe_name(course.code)} - {safe_name(course.name)}"

            for hall_id, hall_data in course_data['halls'].items():
                hall = hall_data['hall']
                cls = hall_data['cls']
                students = hall_data['students']

                # Build DOCX
                doc = Document()

                # Logo
                logo_paragraph = doc.add_paragraph()
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                logo_path = os.path.join(settings.BASE_DIR, 'static', 'assets', 'images', 'logo.png')
                if os.path.exists(logo_path):
                    try:
                        logo_run = logo_paragraph.runs[0] if logo_paragraph.runs else logo_paragraph.add_run()
                        logo_run.add_picture(logo_path, width=Inches(1.0))
                    except Exception:
                        logo_run = logo_paragraph.add_run("[SCHOOL LOGO]")
                        logo_run.bold = True
                else:
                    logo_run = logo_paragraph.add_run("[SCHOOL LOGO]")
                    logo_run.bold = True

                # Session/Semester
                session_paragraph = doc.add_paragraph()
                session_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                session_run = session_paragraph.add_run(f"SESSION: {settings_obj.session}")
                session_run.bold = True
                session_run.add_break()
                semester_run = session_paragraph.add_run(f"SEMESTER: {settings_obj.semester}")
                semester_run.bold = True

                doc.add_paragraph()

                # Course info
                course_header = doc.add_paragraph()
                course_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
                course_run = course_header.add_run(f"COURSE TITLE: {course.name.upper()}")
                course_run.bold = True
                course_run.add_break()
                code_run = course_header.add_run(f"COURSE CODE: {course.code}")
                code_run.bold = True

                # Exam details
                exam_details = doc.add_paragraph()
                exam_details.alignment = WD_ALIGN_PARAGRAPH.LEFT
                hall_run = exam_details.add_run(f"EXAM HALL: {hall.name}")
                hall_run.bold = True
                hall_run.add_break()
                date_run = exam_details.add_run(f"DATE: {date_obj.strftime('%d %B, %Y')}")
                date_run.bold = True

                # Level and period
                level_period = doc.add_paragraph()
                level_period.alignment = WD_ALIGN_PARAGRAPH.LEFT
                level_run = level_period.add_run(f"LEVEL/CLASS: {cls.name}")
                level_run.bold = True
                level_run.add_break()
                period_run = level_period.add_run(f"PERIOD OF EXAM: {period}")
                period_run.bold = True

                # Header
                attendance_header = doc.add_paragraph()
                attendance_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                attendance_run = attendance_header.add_run("ATTENDANCE SHEET")
                attendance_run.bold = True

                doc.add_paragraph()
                doc.add_paragraph()

                # Table
                table = doc.add_table(rows=1, cols=7)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr_cells = table.rows[0].cells
                headers = ['S/NO', 'MATRIC NO', "STUDENT'S NAME", 'SEAT NO', 'SCRIPT NO', 'SIGN IN', 'SIGN OUT']
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    try:
                        hdr_cells[i].paragraphs[0].runs[0].bold = True
                    except IndexError:
                        pass

                for idx, sa in enumerate(students, 1):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(idx)
                    row_cells[1].text = sa.student.matric_no
                    row_cells[2].text = f"{sa.student.first_name} {sa.student.last_name}".upper()
                    row_cells[3].text = str(sa.seat_number) if sa.seat_number else ''

                for i in range(25):
                    row_cells = table.add_row().cells
                    if i < 3:
                        row_cells[0].text = str(len(students) + i + 1)

                doc.add_paragraph()
                doc.add_paragraph()

                footer_info = doc.add_paragraph()
                footer_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
                total_run = footer_info.add_run("TOTAL NUMBER OF STUDENTS……………………………TOTAL NUMBER OF SCRIPTS…………………….")
                total_run.bold = True

                invigilator_info = doc.add_paragraph()
                invigilator_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
                invig_run = invigilator_info.add_run("NAME OF INVIGILATOR SUBMITING SCRIPTS………………………………………………. SIGN…………....")
                invig_run.bold = True

                committee_info = doc.add_paragraph()
                committee_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
                committee_run = committee_info.add_run("NAME OF EXAM COMMITTEE MEMBER RECEIVING SCRIPTS…………………………………………………")
                committee_run.bold = True

                signature_info = doc.add_paragraph()
                signature_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sig_run = signature_info.add_run("SIGNATURE……………………………….                                        DATE…………………………………...")
                sig_run.bold = True

                if students:
                    last_matric = students[-1].student.matric_no.split('/')[-1]
                    range_info = doc.add_paragraph()
                    range_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    range_run = range_info.add_run(f"RANGE: {last_matric}")
                    range_run.bold = True

                # Save document to buffer
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)

                # Path inside zip: course folder + file
                filename = f"Attendance_{safe_name(course.code)}_{safe_name(hall.name)}_{date_obj}_{period}.docx"
                zip_path = f"{course_folder}/{filename}"
                zip_file.writestr(zip_path, doc_buffer.getvalue())

    # Prepare the zip for download
    in_memory_zip.seek(0)
    response = HttpResponse(in_memory_zip.getvalue(), content_type='application/zip')
    response['Content-Disposition'] = f'attachment; filename="Attendance_Sheets_{date_obj}_{period}.zip"'
    return response
