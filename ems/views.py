import io
import os
import zipfile
from datetime import datetime, timedelta
from urllib.parse import urlparse, urlunparse

import pandas as pd
from django.conf import settings
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required, user_passes_test
from django.core.paginator import Paginator
from django.db.models import Count, Prefetch, Q, Sum
from django.http import HttpRequest, HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.views.decorators.http import require_POST
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from .branding import add_document_branding
from .broadsheet import TimetableBroadSheet
from .models import (
    BackgroundJob,
    Class,
    Course,
    Department,
    Distribution,
    DistributionItem,
    Hall,
    SeatArrangement,
    Student,
    SystemSettings,
    TimeTable,
    User,
)
from .utils import handle_uploaded_file


def back_view(request):
    # Get the origin uri
    original_uri = request.META.get("HTTP_REFERER")
    parsed_uri = urlparse(original_uri)

    # Clean up (remove delimiters and empty strings)
    path_parts = parsed_uri.path.split("/")
    path_parts = [path for path in path_parts if path != ""]
    value = path_parts.pop(-1)

    # Construct the referer uri
    new_path_parts = [part for part in path_parts if part != value]
    parsed_uri = parsed_uri._replace(path="/".join(new_path_parts))

    referer = urlunparse(parsed_uri)

    # redirect to the referer uri
    if referer:
        return redirect(referer)
    else:
        return redirect("dashboard/")


def admin_required(view_func):
    decorated_view_func = user_passes_test(lambda user: user.is_staff)(view_func)
    return decorated_view_func


def index(request):
    return render(request, template_name="site/index.html")

def features_view(request):
    return render(request, template_name="site/features.html")

def feature_detail_view(request, slug):
    features_data = {
        'timetable-generation': {
            'title': 'Timetable Generation',
            'subtitle': 'Automatically create conflict-free exam schedules.',
            'overview': 'The timetable generation module is the core of ExamNova. It automatically builds a comprehensive examination schedule based on courses, classes, and available time slots. It ensures that no class has conflicting exams, appropriately handles different exam types (PBE vs CBE), and intelligently spaces out the schedule.',
            'icon_class': 'indigo',
            'icon': 'fas fa-calendar-check',
            'capabilities': [
                'AM/PM period scheduling',
                'CBE vs PBE exam type handling (exclusive slots for CBE)',
                'Sunday exclusion & date range validation',
                'Background processing with real-time progress tracking',
                'Department-level and institution-wide views',
                'Export to CSV and comprehensive Excel broadsheet'
            ],
            'how_it_works': [
                'Upload all required data (departments, courses, classes, halls).',
                'Set the start and end dates for the examination period.',
                'Click "Generate" and monitor the progress.',
                'Review and export the generated timetable.'
            ],
            'benefits': [
                'Saves days of manual scheduling work.',
                'Eliminates human errors and exam conflicts.',
                'Provides a clear, organized schedule for both staff and students.'
            ]
        },
        'hall-distribution': {
            'title': 'Hall Distribution',
            'subtitle': 'Intelligently distribute class groups across available examination halls.',
            'overview': 'Once the timetable is set, the Hall Distribution module takes over. It determines which classes will sit in which halls for each specific date and period. It optimizes for hall capacity, ensuring halls are efficiently utilized without overcrowding.',
            'icon_class': 'cyan',
            'icon': 'fas fa-th-large',
            'capabilities': [
                'Capacity-aware hall assignment',
                'Large class splitting across multiple halls',
                'Utilization statistics & optimization grade',
                'Distribution CSV export',
                'Per-date, per-period configuration'
            ],
            'how_it_works': [
                'Select a specific date and period from the generated timetable.',
                'Click "Generate Distribution".',
                'Review the assigned halls and class groups.',
                'Check the efficiency statistics to ensure optimal usage.'
            ],
            'benefits': [
                'Maximizes space utilization.',
                'Prevents overcrowding in examination halls.',
                'Provides clear metrics on how well resources are being used.'
            ]
        },
        'seat-allocation': {
            'title': 'Seat Allocation',
            'subtitle': 'Assign specific seats to students with strict anti-cheating constraints.',
            'overview': 'The Seat Allocation module is an advanced seating algorithm designed to prevent cheating. It assigns specific seat numbers to individual students, ensuring that no two students taking the same course are seated next to each other in any direction.',
            'icon_class': 'emerald',
            'icon': 'fas fa-grip-vertical',
            'capabilities': [
                '8-directional adjacency enforcement (horizontal, vertical, diagonal)',
                'Pattern-based seating (checkerboard, diagonal multi-pass)',
                'Visual seat grid layout per hall',
                'Manual seat assignment for unplaced students',
                'Detailed placed vs unplaced summary'
            ],
            'how_it_works': [
                'After distribution, generate seat allocation for a date and period.',
                'The system processes each hall, applying adjacency constraints.',
                'Review the visual seat grid for each hall.',
                'Manually assign seats to any unplaced students if necessary.'
            ],
            'benefits': [
                'Significantly reduces the potential for cheating.',
                'Provides a clear, organized seating plan for invigilators.',
                'Visual grid makes it easy to understand the hall layout.'
            ]
        },
        'reports-exports': {
            'title': 'Reports & Exports',
            'subtitle': 'Generate professional documents for exam halls.',
            'overview': 'ExamNova provides a comprehensive suite of reporting and export tools to generate all the necessary documentation for conducting examinations.',
            'icon_class': 'indigo',
            'icon': 'fas fa-file-word',
            'capabilities': [
                'DOCX attendance sheets with school branding and signature fields',
                'Excel broadsheet (department & day-period views)',
                'CSV timetable export per department',
                'ZIP packages of seat arrangements per course',
                'Walk-in rows for late registrations'
            ],
            'benefits': [
                'Produces ready-to-print, professional documents.',
                'Reduces manual paperwork.',
                'Ensures consistency across all examination records.'
            ]
        },
        'data-management': {
            'title': 'Data Management',
            'subtitle': 'Complete CRUD and bulk CSV upload support for all system entities.',
            'overview': 'A robust data management system allows administrators and exam officers to easily input and manage the foundational data required for the examination process.',
            'icon_class': 'cyan',
            'icon': 'fas fa-database',
            'capabilities': [
                'Manual create, edit, delete for all entities',
                'Bulk CSV upload with strict validation',
                'ZIP-based multi-file upload support (Bulk Upload Module)',
                'Duplicate detection and conflict handling',
                'Upload lock after timetable generation to maintain data integrity'
            ],
            'benefits': [
                'Fast and efficient data entry via bulk uploads.',
                'Maintains data integrity with strict validation rules.',
                'Clear error reporting helps quickly identify data issues.'
            ]
        },
        'background-jobs': {
            'title': 'Background Job Monitor',
            'subtitle': 'Track, retry, and manage all long-running tasks.',
            'overview': 'ExamNova uses asynchronous background processing for heavy tasks like timetable generation and seat allocation. The Job Monitor provides full visibility into these processes.',
            'icon_class': 'emerald',
            'icon': 'fas fa-tasks',
            'capabilities': [
                'Real-time progress tracking',
                'Filter by status (Pending, Running, Success, Failed) and job type',
                'Retry failed jobs with one click',
                'Detailed error messages and tracebacks for debugging',
                'Job history with parameters and result metrics'
            ],
            'benefits': [
                'Ensures the system remains responsive during heavy processing.',
                'Provides transparency into the status of complex operations.',
                'Easy recovery from failures via the retry mechanism.'
            ]
        }
    }
    
    feature = features_data.get(slug)
    if not feature:
        # Redirect to features overview if slug is invalid
        return redirect('features')
        
    return render(request, template_name="site/feature-detail.html", context={'feature': feature})



def login_view(request):
    if request.user.is_authenticated:
        return redirect("dashboard")
    if request.method == "POST":
        email = request.POST.get("email")
        password = request.POST.get("password")
        user = authenticate(request, username=email, password=password)
        if user is not None:
            login(request, user)
            return redirect("dashboard")
        messages.error(request, "Invalid email or password.")
        return redirect("dashboard")

    return render(
        request,
        template_name="site/login.html",
    )


@login_required(login_url="login")
def logout_view(request):
    logout(request)
    return redirect("login")


@login_required(login_url="login")
def dashboard(request):
    settings = SystemSettings.objects.first()
    if not settings:
        SystemSettings.objects.create(session="2024/2025", semester="1st Semester")

    # Filter statistics based on user role
    if request.user.is_staff:
        # Admin users see all data
        departments = Department.objects.all()
        departments_count = departments.count()
        halls = Hall.objects.all().count()
        courses_count = Course.objects.all().count()
        classes_count = Class.objects.all().count()
        students = Class.objects.aggregate(total_size=Sum("size"))["total_size"] or 0

        # Courses shared across multiple departments (optimized)
        shared_courses_qs = (
            Course.objects.annotate(
                dept_count=Count("courses__department", distinct=True)
            )
            .filter(dept_count__gt=1)
            .prefetch_related(
                Prefetch(
                    "courses",
                    queryset=Class.objects.select_related("department").only(
                        "id", "name", "department__name", "department_id"
                    ),
                    to_attr="prefetched_classes",
                )
            )
            .order_by("-dept_count")
        )

        shared_courses = []
        for course in shared_courses_qs:
            dept_classes_map = {}
            for cls in course.prefetched_classes:
                dept_name = cls.department.name
                dept_classes_map.setdefault(dept_name, []).append(cls.name)

            dept_with_classes = [
                {"name": dept_name, "classes": class_names}
                for dept_name, class_names in dept_classes_map.items()
            ]

            shared_courses.append(
                {
                    "code": course.code,
                    "name": course.name,
                    "dept_count": course.dept_count,
                    "departments": dept_with_classes,
                }
            )
    else:
        # Non-admin users see only their department data
        if request.user.department:
            departments_count = 1  # Only their department
            halls = Hall.objects.all().count()  # Halls are shared resources
            # Courses from classes in their department
            dept_classes = Class.objects.filter(department=request.user.department)
            courses_count = (
                Course.objects.filter(courses__in=dept_classes).distinct().count()
            )
            classes_count = dept_classes.count()
            students = dept_classes.aggregate(total_size=Sum("size"))["total_size"] or 0
            # Non-admin users don't see shared courses (single department view)
            shared_courses = []
        else:
            departments_count = courses_count = students = classes_count = 0
            halls = 0
            shared_courses = []

    context = {
        "departments_count": departments_count,
        "halls_count": halls,
        "courses_count": courses_count,
        "classes_count": classes_count,
        "students": students,
        "shared_courses": shared_courses,
        "shared_courses_count": len(shared_courses),
        "settings": settings,
    }
    if request.htmx:
        template_name = "dashboard/pages/dashboard.html"
    else:
        template_name = "dashboard/index.html"

    return render(request, template_name=template_name, context=context)


@login_required(login_url="login")
@admin_required
def setting(request):
    settings = SystemSettings.objects.first()
    if not settings:
        settings = SystemSettings.objects.create(
            session="2024/2025", semester="1st Semester"
        )
    context = {"settings": settings}
    if request.htmx:
        template_name = "dashboard/pages/settings.html"
    else:
        template_name = "dashboard/settings.html"

    return render(request, template_name=template_name, context=context)


@require_POST
@login_required(login_url="login")
@admin_required
def update_settings(request):
    session = request.POST.get("session")
    semester = request.POST.get("semester")

    settings, created = SystemSettings.objects.get_or_create(id=1)

    settings.session = session
    settings.semester = semester
    settings.is_generated = True
    settings.save()

    if created:
        messages.success(request, "System settings created successfully.")
    else:
        messages.success(request, "System settings updated successfully.")

    return redirect("settings")


@require_POST
@login_required(login_url="login")
@admin_required
def manual_seat_assignment(request):
    """
    Manually assign an unplaced student to an available seat
    """
    student_id = request.POST.get("student_id")
    seat_number = request.POST.get("seat_number")
    date = request.POST.get("date")
    period = request.POST.get("period")
    hall_id = request.POST.get("hall_id")

    if not all([student_id, seat_number, date, period, hall_id]):
        messages.error(request, "Missing required parameters for manual assignment.")
        return redirect(
            f"/hall-allocation/?date={date}&period={period}&hall_id={hall_id}"
        )

    try:
        # Get the unplaced student
        unplaced_student = SeatArrangement.objects.get(
            id=student_id,
            date=date,
            period=period,
            hall_id=hall_id,
            seat_number__isnull=True,
        )

        # Check if the seat is available
        existing_assignment = SeatArrangement.objects.filter(
            date=date, period=period, hall_id=hall_id, seat_number=seat_number
        ).first()

        if existing_assignment:
            messages.error(
                request,
                f"Seat {seat_number} is already occupied by {existing_assignment.student.first_name} {existing_assignment.student.last_name}.",
            )
            return redirect(
                f"/hall-allocation/?date={date}&period={period}&hall_id={hall_id}"
            )

        # Get hall to validate seat number
        hall = Hall.objects.get(id=hall_id)
        max_seats = hall.rows * hall.columns if hall.rows and hall.columns else 0

        if int(seat_number) > max_seats or int(seat_number) < 1:
            messages.error(
                request, f"Invalid seat number. Hall capacity is {max_seats} seats."
            )
            return redirect(
                f"/hall-allocation/?date={date}&period={period}&hall_id={hall_id}"
            )

        # Skip adjacency constraints for manual assignment to allow flexible placement

        # Assign the seat
        unplaced_student.seat_number = int(seat_number)
        unplaced_student.save()

        messages.success(
            request,
            f"Successfully assigned {unplaced_student.student.first_name} {unplaced_student.student.last_name} to seat {seat_number}.",
        )

    except SeatArrangement.DoesNotExist:
        messages.error(request, "Student not found or already placed.")
    except Hall.DoesNotExist:
        messages.error(request, "Hall not found.")
    except Exception as e:
        messages.error(request, f"Error during manual assignment: {str(e)}")

    return redirect(f"/hall-allocation/?date={date}&period={period}&hall_id={hall_id}")


@login_required(login_url="login")
def departments(request):
    # Filter departments based on user role
    if request.user.is_staff:
        departments_list = Department.objects.all().order_by("id")
    else:
        # Non-admin users can only see their own department
        departments_list = (
            Department.objects.filter(id=request.user.department.id)
            if request.user.department
            else Department.objects.none()
        )

    query = request.GET.get("query")
    if query:
        departments_list = departments_list.filter(
            Q(name__icontains=query) | Q(slug__icontains=query)
        )

    paginator = Paginator(departments_list, 15)
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)
    context = {"departments": page_obj}
    if request.htmx:
        template_name = "dashboard/pages/departments.html"
    else:
        template_name = "dashboard/departments.html"

    return render(request, template_name=template_name, context=context)


@login_required(login_url="login")
def get_department(request, slug):
    department = get_object_or_404(Department, slug=slug)

    # Check if user has permission to view this department
    if not request.user.is_staff and request.user.department != department:
        messages.error(request, "You don't have permission to view this department.")
        return redirect("departments")

    classes = Class.objects.filter(department=department)
    context = {"department": department, "classes": classes}
    if request.htmx:
        template_name = "dashboard/pages/single-department.html"
    else:
        template_name = "dashboard/single-department.html"

    return render(request, template_name=template_name, context=context)


@login_required(login_url="login")
def get_courses_view(request):
    # Filter courses based on user role
    if request.user.is_staff:
        course_list = Course.objects.all().order_by("id")
    else:
        # Non-admin users see only courses from their department's classes
        if request.user.department:
            dept_classes = Class.objects.filter(department=request.user.department)
            course_list = (
                Course.objects.filter(courses__in=dept_classes)
                .distinct()
                .order_by("id")
            )
        else:
            course_list = Course.objects.none()

    query = request.GET.get("query")
    if query:
        course_list = course_list.filter(
            Q(name__icontains=query) | Q(code__icontains=query)
        )

    paginator = Paginator(course_list, 15)
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)
    context = {"courses": page_obj}
    if request.htmx:
        template_name = "dashboard/pages/courses.html"
    else:
        template_name = "dashboard/courses.html"

    return render(request, template_name=template_name, context=context)


@login_required(login_url="login")
@admin_required
def get_students(request):
    students = Student.objects.all().order_by("id")
    query = request.GET.get("query")
    if query:
        students = students.filter(
            Q(first_name__icontains=query)
            | Q(last_name__icontains=query)
            | Q(matric_no__icontains=query)
            | Q(email__icontains=query)
            | Q(phone__icontains=query)
        )

    paginator = Paginator(students, 15)
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)
    context = {"students": page_obj}
    if request.htmx:
        template_name = "dashboard/pages/students.html"
    else:
        template_name = "dashboard/students.html"

    return render(request, template_name=template_name, context=context)


@login_required(login_url="login")
def get_class_course(request, slug, id):
    cls = get_object_or_404(Class, department__slug=slug, id=id)

    # Check if user has permission to view this class
    if not request.user.is_staff and request.user.department != cls.department:
        messages.error(request, "You don't have permission to view this class.")
        return redirect("departments")

    students = Student.objects.filter(level=cls, department=cls.department)
    context = {"class": cls, "courses": cls.courses.all(), "students": students}
    if request.htmx:
        template_name = "dashboard/pages/class.html"
    else:
        template_name = "dashboard/class.html"

    return render(request, template_name=template_name, context=context)


@login_required(login_url="login")
@admin_required
def halls(request):
    halls_list = Hall.objects.all()
    paginator = Paginator(halls_list, 15)
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)
    context = {"halls": page_obj}
    if request.htmx:
        template_name = "dashboard/pages/halls.html"
    else:
        template_name = "dashboard/halls.html"

    return render(request, template_name=template_name, context=context)


@login_required(login_url="login")
def timetable(request: HttpRequest) -> HttpResponse:
    generated = TimeTable.objects.exists()
    context = {"generated": generated}
    if generated:
        dates = (
            TimeTable.objects.values_list("date", flat=True).distinct().order_by("date")
        )
        date = request.GET.get("date")
        period = request.GET.get("period")
        if date is not None or period is not None:
            timetables = TimeTable.objects.filter(date=date, period=period)
        else:
            timetables = TimeTable.objects.filter(date=dates[0], period="AM")
        if request.user.is_staff is False:
            timetables = timetables.filter(
                class_obj__department=request.user.department
            )
        context["dates"] = dates
        context["timetables"] = timetables

    if request.htmx:
        template_name = "dashboard/pages/timetable.html"
    else:
        template_name = "dashboard/timetable.html"

    return render(request, template_name=template_name, context=context)


@login_required(login_url="login")
def distribution(request):
    generated = Distribution.objects.exists()
    dates = TimeTable.objects.values_list("date", flat=True).distinct().order_by("date")
    date = request.GET.get("date")
    period = request.GET.get("period")

    # Default to first date and AM if not provided
    if not date or not period:
        first_date = dates.first() if dates.exists() else None
        if not date:
            date = str(first_date) if first_date else None
        if not period:
            period = "AM"

    context = {"generated": generated, "dates": dates, "date": date, "period": period}
    if generated:
        if date and period:
            distributions = Distribution.objects.filter(date=date, period=period)
        else:
            distributions = Distribution.objects.none()

        # Filter distributions based on user role
        if not request.user.is_staff and request.user.department:
            # Non-admin users see only distributions for their department's classes
            distributions = distributions.filter(
                items__schedule__class_obj__department=request.user.department
            ).distinct()

        context["distributions"] = distributions

    if request.htmx:
        template_name = "dashboard/pages/distribution.html"
    else:
        template_name = "dashboard/distribution.html"

    return render(request, template_name=template_name, context=context)


@login_required(login_url="login")
def allocation(request):
    generated = SeatArrangement.objects.exists()
    dates = TimeTable.objects.values_list("date", flat=True).distinct().order_by("date")
    print(dates)
    date = request.GET.get("date")
    period = request.GET.get("period")
    hall_id = request.GET.get("hall")

    context = {
        "generated": generated,
        "dates": dates,
        "date": date,
        "period": period,
        "hall_id": hall_id,
    }

    if generated:
        first_date = dates.first() if dates.exists() else None
        # Determine which date and period to use
        if date and period:  # Check if both date and period are not empty
            try:
                # Validate date format
                from datetime import datetime

                datetime.strptime(date, "%Y-%m-%d")
                arrangements = SeatArrangement.objects.filter(date=date, period=period)
            except (ValueError, TypeError):
                # If date is invalid, use default
                if first_date:
                    arrangements = SeatArrangement.objects.filter(
                        date=first_date, period="AM"
                    )
                    date = first_date
                    period = "AM"
                else:
                    arrangements = SeatArrangement.objects.none()
        else:
            if first_date:
                arrangements = SeatArrangement.objects.filter(
                    date=first_date, period="AM"
                )
                date = first_date
                period = "AM"
            else:
                arrangements = SeatArrangement.objects.none()

        # Filter arrangements based on user role
        if not request.user.is_staff and request.user.department:
            # Non-admin users see only arrangements for their department's classes
            arrangements = arrangements.filter(cls__department=request.user.department)

        hall_arrangements = arrangements.values(
            "hall__name", "hall__id", "date", "period"
        ).annotate(
            placed=Count("id", filter=Q(seat_number__isnull=False)),
            not_placed=Count("id", filter=Q(seat_number__isnull=True)),
        )
        context["arrangements"] = hall_arrangements

    if request.htmx:
        template_name = "dashboard/pages/allocation.html"
    else:
        template_name = "dashboard/allocation.html"

    return render(request, template_name=template_name, context=context)


@login_required(login_url="login")
def hall_allocation(request):
    """
    View for detailed hall allocation with visual seat layout
    """
    generated = SeatArrangement.objects.exists()
    dates = TimeTable.objects.values_list("date", flat=True).distinct().order_by("date")
    date = request.GET.get("date")
    period = request.GET.get("period")
    hall_id = request.GET.get("hall_id")

    if not hall_id:
        messages.error(request, "Hall ID is required for detailed view.")
        return redirect("allocation")

    context = {
        "generated": generated,
        "dates": dates,
        "date": date,
        "period": period,
        "hall_id": hall_id,
    }

    if generated:
        # Determine which date and period to use
        if date and period:
            try:
                from datetime import datetime

                datetime.strptime(date, "%Y-%m-%d")
                arrangements = SeatArrangement.objects.filter(
                    date=date, period=period, hall_id=hall_id
                )
            except (ValueError, TypeError):
                arrangements = SeatArrangement.objects.filter(
                    date=dates[0], period="AM", hall_id=hall_id
                )
                date = dates[0]
                period = "AM"
        else:
            arrangements = SeatArrangement.objects.filter(
                date=dates[0], period="AM", hall_id=hall_id
            )

        # Filter arrangements based on user role
        if not request.user.is_staff and request.user.department:
            # Non-admin users see only arrangements for their department's classes
            arrangements = arrangements.filter(cls__department=request.user.department)

        # Get detailed seat arrangement for the specific hall
        hall_details = arrangements.select_related(
            "student", "course", "cls", "hall"
        ).order_by("seat_number")

        # Separate placed and unplaced students
        placed_students = hall_details.filter(seat_number__isnull=False)
        unplaced_students = hall_details.filter(seat_number__isnull=True)

        # Get hall info for seating grid
        if hall_details.exists():
            hall = hall_details.first().hall

            # Create a 2D grid for visual representation
            seat_grid = []
            if hall.rows and hall.columns:
                # Initialize empty grid
                for row in range(hall.rows):
                    seat_row = []
                    for col in range(hall.columns):
                        seat_number = row * hall.columns + col + 1
                        seat_row.append(
                            {
                                "seat_number": seat_number,
                                "student": None,
                                "course": None,
                                "is_occupied": False,
                            }
                        )
                    seat_grid.append(seat_row)

                # Fill grid with placed students
                for student in placed_students:
                    if student.seat_number:
                        # Convert seat number to row, col (1-indexed to 0-indexed)
                        seat_num = student.seat_number - 1
                        row = seat_num // hall.columns
                        col = seat_num % hall.columns

                        if 0 <= row < hall.rows and 0 <= col < hall.columns:
                            seat_grid[row][col].update(
                                {
                                    "student": student.student,
                                    "course": student.course,
                                    "is_occupied": True,
                                    "student_matric_no": getattr(
                                        student, "student_matric_no", ""
                                    ),
                                    "cls": student.cls,
                                }
                            )

            context.update(
                {
                    "hall_details": hall_details,
                    "placed_students": placed_students,
                    "unplaced_students": unplaced_students,
                    "selected_hall": hall,
                    "hall_capacity": hall.rows * hall.columns
                    if hall.rows and hall.columns
                    else 0,
                    "seat_grid": seat_grid,
                }
            )

    if request.htmx:
        template_name = "dashboard/pages/hall-allocation.html"
    else:
        template_name = "dashboard/hall-allocation.html"

    return render(request, template_name=template_name, context=context)


@login_required(login_url="login")
@admin_required
def generate_attendance_sheets(request):
    """Generate attendance sheets for all courses in a specific hall"""
    date = request.GET.get("date")
    period = request.GET.get("period")
    hall_id = request.GET.get("hall_id")

    if not all([date, period, hall_id]):
        messages.error(request, "Date, period, and hall must be specified.")
        return redirect("hall-allocation")

    try:
        # Validate date format
        datetime.strptime(date, "%Y-%m-%d")
    except ValueError:
        messages.error(request, "Invalid date format.")
        return redirect("allocation")

    # Get system settings for session and semester
    settings_obj = SystemSettings.objects.first()
    if not settings_obj:
        settings_obj = SystemSettings.objects.create(
            session="2024/2025", semester="1st Semester"
        )

    # Get arrangements for the specific hall, date, and period
    arrangements = (
        SeatArrangement.objects.filter(date=date, period=period, hall_id=hall_id)
        .select_related("student", "course", "cls", "hall")
        .order_by("course__name", "student__matric_no")
    )

    if not arrangements.exists():
        messages.error(
            request, "No seat arrangements found for the specified criteria."
        )
        return redirect("allocation")

    hall = arrangements.first().hall

    # Group students by course
    courses_data = {}
    for arrangement in arrangements:
        if arrangement.seat_number:  # Only placed students
            course_key = f"{arrangement.course.name} ({arrangement.course.code})"
            if course_key not in courses_data:
                courses_data[course_key] = {
                    "course": arrangement.course,
                    "cls": arrangement.cls,
                    "students": [],
                }
            courses_data[course_key]["students"].append(arrangement)

    if not courses_data:
        messages.error(request, "No placed students found for attendance sheets.")
        return redirect("allocation")

    # Create a zip file containing all attendance sheets
    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for course_key, course_data in courses_data.items():
            # Create Word document for each course
            doc = Document()

            # Institution branding header (logo + metadata + session).
            add_document_branding(doc, settings_obj)

            # Add course information header
            course_header = doc.add_paragraph()
            course_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
            course_run = course_header.add_run(
                f"COURSE TITLE: {course_data['course'].name.upper()}"
            )
            course_run.bold = True
            course_run.add_break()
            code_run = course_header.add_run(
                f"COURSE CODE: {course_data['course'].code}"
            )
            code_run.bold = True

            # Add exam details
            exam_details = doc.add_paragraph()
            exam_details.alignment = WD_ALIGN_PARAGRAPH.LEFT
            hall_run = exam_details.add_run(f"EXAM HALL: {hall.name}")
            hall_run.bold = True
            hall_run.add_break()

            date_run = exam_details.add_run(
                f"DATE: {datetime.strptime(date, '%Y-%m-%d').strftime('%d %B, %Y')}"
            )
            date_run.bold = True

            # Add level and period
            level_period = doc.add_paragraph()
            level_period.alignment = WD_ALIGN_PARAGRAPH.LEFT
            level_run = level_period.add_run(
                f"LEVEL/CLASS: {course_data['cls'].full_label}"
            )
            level_run.bold = True
            level_run.add_break()
            period_run = level_period.add_run(f"PERIOD OF EXAM: {period}")
            period_run.bold = True

            # Add attendance sheet header
            attendance_header = doc.add_paragraph()
            attendance_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            attendance_run = attendance_header.add_run("ATTENDANCE SHEET")
            attendance_run.bold = True

            # Add spacing
            doc.add_paragraph()
            doc.add_paragraph()

            # Create attendance table
            table = doc.add_table(rows=1, cols=7)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Add table headers
            hdr_cells = table.rows[0].cells
            headers = [
                "S/NO",
                "MATRIC NO",
                "STUDENT'S NAME",
                "SEAT NO",
                "SCRIPT NO",
                "SIGN IN",
                "SIGN OUT",
            ]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].bold = True

            # Add student data
            for idx, student in enumerate(course_data["students"], 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = student.student.matric_no
                row_cells[
                    2
                ].text = (
                    f"{student.student.first_name} {student.student.last_name}".upper()
                )
                row_cells[3].text = (
                    str(student.seat_number) if student.seat_number else ""
                )
                # Leave script no, sign in, sign out empty for manual filling

            # Add extra blank rows (20-25 as requested)
            for i in range(25):
                row_cells = table.add_row().cells
                if i < 3:  # First 3 extra rows have numbers
                    row_cells[0].text = str(len(course_data["students"]) + i + 1)

            # Add spacing
            doc.add_paragraph()
            doc.add_paragraph()

            # Add footer information
            footer_info = doc.add_paragraph()
            footer_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            total_run = footer_info.add_run(
                "TOTAL NUMBER OF STUDENTS……………………………TOTAL NUMBER OF SCRIPTS……………………."
            )
            total_run.bold = True

            invigilator_info = doc.add_paragraph()
            invigilator_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            invig_run = invigilator_info.add_run(
                "NAME OF INVIGILATOR SUBMITING SCRIPTS………………………………………………. SIGN…………...."
            )
            invig_run.bold = True

            committee_info = doc.add_paragraph()
            committee_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            committee_run = committee_info.add_run(
                "NAME OF EXAM COMMITTEE MEMBER RECEIVING SCRIPTS…………………………………………………"
            )
            committee_run.bold = True

            signature_info = doc.add_paragraph()
            signature_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sig_run = signature_info.add_run(
                "SIGNATURE……………………………….                                        DATE…………………………………..."
            )
            sig_run.bold = True

            # Add range information
            if course_data["students"]:
                last_matric = course_data["students"][-1].student.matric_no.split("/")[
                    -1
                ]
                range_info = doc.add_paragraph()
                range_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
                range_run = range_info.add_run(f"RANGE: {last_matric}")
                range_run.bold = True

            # Save document to buffer
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            # Add to zip file
            filename = f"Attendance_{course_data['course'].code}_{hall.name}_{date}_{period}.docx"
            zip_file.writestr(filename, doc_buffer.getvalue())

    # Prepare response
    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer.getvalue(), content_type="application/zip")
    response["Content-Disposition"] = (
        f'attachment; filename="Attendance_Sheets_{hall.name.strip()}_{date.strip()}_{period.strip()}.zip"'
    )

    return response


@login_required(login_url="login")
@admin_required
def generate_broadsheet(request):
    """Export timetable as Excel file"""

    # To be collected system setting
    settings = SystemSettings.objects.first()
    semester = settings.semester
    academic_year = settings.session

    try:
        # Get all timetables
        timetables = TimeTable.objects.select_related(
            "course", "class_obj", "class_obj__department"
        ).all()

        # Order by date and period
        timetables = timetables.order_by("date", "period")

        if not timetables.exists():
            return HttpResponse(
                "No examination records found for the specified criteria.", status=404
            )

        # Generate Excel file
        exporter = TimetableBroadSheet()
        workbook = exporter.generate_excel(list(timetables), semester, academic_year)

        # Create HTTP response
        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        # Generate filename
        filename = f"Examination_Timetable_{semester}_Semester_{academic_year.replace('/', '_')}.xlsx"
        response["Content-Disposition"] = f'attachment; filename="{filename}"'

        # Save workbook to response
        virtual_workbook = io.BytesIO()
        workbook.save(virtual_workbook)
        virtual_workbook.seek(0)
        response.write(virtual_workbook.getvalue())

        return response

    except Exception as e:
        return HttpResponse(f"Error generating Excel file: {str(e)}", status=500)


@login_required(login_url="login")
@admin_required
def manage_users(request):
    users = User.objects.all()
    department_list = Department.objects.all()
    query = request.GET.get("query")
    if query:
        users = users.filter(
            Q(first_name__icontains=query) | Q(last_name__icontains=query)
        )

    paginator = Paginator(users, 10)
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)
    context = {"users": page_obj, "departments": department_list}

    if request.htmx:
        template_name = "dashboard/pages/manage-users.html"
    else:
        template_name = "dashboard/manage-users.html"

    return render(request, template_name=template_name, context=context)


@require_POST
@login_required(login_url="login")
@admin_required
def add_user(request):
    first_name = request.POST.get("first_name")
    last_name = request.POST.get("last_name")
    email = request.POST.get("email")
    department_slug = request.POST.get("department")
    password = request.POST.get("password")
    password_confirm = request.POST.get("password-confirm")

    if password != password_confirm:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "Passwords do not match"},
        )
    if not Department.objects.filter(slug=department_slug).exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": f"Department with code {department_slug} does not exists"
            },
        )
    if User.objects.filter(email=email).exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "Email already exists"},
        )

    user = User.objects.create_user(
        first_name=first_name,
        last_name=last_name,
        email=email,
        department=get_object_or_404(Department, slug=department_slug),
        password=password,
    )
    user.save()
    return render(
        request,
        template_name="dashboard/partials/alert-success.html",
        context={"message": "User created successfully"},
    )


@login_required(login_url="login")
@admin_required
def reset_system(request: HttpRequest) -> HttpResponse:
    SeatArrangement.objects.all().delete()
    Distribution.objects.all().delete()
    DistributionItem.objects.all().delete()
    TimeTable.objects.all().delete()
    Hall.objects.all().delete()
    Course.objects.all().delete()
    Class.objects.all().delete()
    Student.objects.all().delete()
    Department.objects.all().delete()

    setting = SystemSettings.objects.first()
    setting.has_timetable = False
    setting.save()

    return redirect("dashboard")


# --------------------------
#  Generate Views
# --------------------------


@require_POST
@login_required(login_url="login")
@admin_required
def generate_timetable(request: HttpRequest) -> HttpResponse:
    import uuid

    from .models import BackgroundJob
    from .tasks import generate_timetable_task

    startDate = request.POST.get("startDate")
    endDate = request.POST.get("endDate")

    if startDate == "" or endDate == "":
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "Please select a start date and end date"},
        )

    # Validation 1: Check if any courses exist in the system
    if not Course.objects.exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": "Cannot generate timetable. No courses found in the system. Please upload courses first."
            },
        )

    # Validation 2: Check if any classes exist in the system
    if not Class.objects.exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": "Cannot generate timetable. No classes found in the system. Please upload classes first."
            },
        )

    # Validation 3: Check if any halls exist in the system
    if not Hall.objects.exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": "Cannot generate timetable. No halls found in the system. Please upload halls first."
            },
        )

    # Validation 4: Check if all classes have at least one course assigned
    classes_without_courses = Class.objects.filter(courses__isnull=True).select_related(
        "department"
    )
    if classes_without_courses.exists():
        table_data = [
            [cls.department.name, cls.name] for cls in classes_without_courses
        ]
        return render(
            request,
            template_name="dashboard/partials/alert-error-table.html",
            context={
                "title": "Cannot Generate Timetable",
                "description": "The following classes have no courses assigned. Please upload courses for these classes first.",
                "table_headers": ["Department", "Class"],
                "table_data": table_data,
            },
        )

    startDate_obj = datetime.strptime(startDate, "%Y-%m-%d").date()
    endDate_obj = datetime.strptime(endDate, "%Y-%m-%d").date()

    if startDate_obj > endDate_obj:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "End date must be greater than start date"},
        )

    dates = []
    currentDate = startDate_obj
    while currentDate <= endDate_obj:
        if currentDate.weekday() != 6:  # 6 represents Sunday
            dates.append(currentDate)
        currentDate += timedelta(days=1)

    # Validation 5: Check if selected date range provides enough days for timetable generation
    # Find the class with the highest number of courses
    # This determines the minimum days needed since each course requires one exam slot
    max_courses_per_class = 0
    for cls in Class.objects.prefetch_related("courses"):
        course_count = cls.courses.count()
        if course_count > max_courses_per_class:
            max_courses_per_class = course_count

    min_days_needed = max_courses_per_class
    available_days = len(dates)

    if available_days < min_days_needed:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": f"Cannot generate timetable. Selected date range provides {available_days} days but minimum {min_days_needed} days are required (based on class with most courses). Please select a longer date range."
            },
        )

    # Create background job
    job_id = str(uuid.uuid4())
    job = BackgroundJob.objects.create(
        job_id=job_id,
        job_type="timetable",
        status="pending",
        created_by=request.user,
        params={"start_date": startDate, "end_date": endDate},
    )

    # Force commit to database before triggering task (important for SQLite)
    from django.db import transaction

    transaction.on_commit(
        lambda: generate_timetable_task.apply_async(
            args=[job_id, request.user.id, startDate, endDate]
        )
    )

    return render(
        request,
        template_name="dashboard/partials/job-started.html",
        context={"job_id": job_id, "job_type": "Timetable Generation"},
    )


@require_POST
@login_required(login_url="login")
@admin_required
def generate_distribution(request: HttpRequest) -> HttpResponse:
    import uuid

    from .models import BackgroundJob
    from .tasks import generate_distribution_task

    date = request.POST.get("date")
    period = request.POST.get("period")

    if not Distribution.objects.filter(date=date, period=period).exists():
        # Create background job
        job_id = str(uuid.uuid4())
        job = BackgroundJob.objects.create(
            job_id=job_id,
            job_type="distribution",
            status="pending",
            created_by=request.user,
            params={"date": date, "period": period},
        )

        # Force commit to database before triggering task (important for SQLite)
        from django.db import transaction

        transaction.on_commit(
            lambda: generate_distribution_task.apply_async(
                args=[job_id, request.user.id, date, period]
            )
        )

        return render(
            request,
            template_name="dashboard/partials/job-started.html",
            context={"job_id": job_id, "job_type": "Distribution Generation"},
        )
    else:
        # Distribution already exists — show inline message and refresh the list
        response = render(
            request,
            template_name="dashboard/partials/distribution-exists.html",
            context={"date": date, "period": period},
        )
        response["HX-Trigger"] = "distributionReady"
        return response


@require_POST
@login_required(login_url="login")
@admin_required
def generate_allocation(request: HttpRequest) -> HttpResponse:
    """
    Generate optimized seat allocation with multi-pass strategy and flexible constraints.
    """
    import uuid

    from .models import BackgroundJob
    from .tasks import generate_allocation_task

    date = request.POST.get("date")
    period = request.POST.get("period")

    if not SeatArrangement.objects.filter(date=date, period=period).exists():
        distributions = Distribution.objects.filter(date=date, period=period)

        if not distributions.exists():
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={
                    "message": f"No distribution found for {date} period {period}. Please generate distribution first."
                },
            )

        # Create background job
        job_id = str(uuid.uuid4())
        job = BackgroundJob.objects.create(
            job_id=job_id,
            job_type="allocation",
            status="pending",
            created_by=request.user,
            params={"date": date, "period": period},
        )

        # Force commit to database before triggering task (important for SQLite)
        from django.db import transaction

        transaction.on_commit(
            lambda: generate_allocation_task.apply_async(
                args=[job_id, request.user.id, date, period]
            )
        )

        return render(
            request,
            template_name="dashboard/partials/job-started.html",
            context={"job_id": job_id, "job_type": "Seat Allocation"},
        )
    else:
        # Allocation already exists - redirect to view it
        from django.http import HttpResponse

        return HttpResponse(
            headers={
                "HX-Redirect": reverse("allocation") + f"?date={date}&period={period}"
            }
        )


@login_required(login_url="login")
def check_job_status(request: HttpRequest, job_id: str) -> HttpResponse:
    """
    API endpoint to check the status of a background job
    """
    from .models import BackgroundJob

    try:
        job = BackgroundJob.objects.get(job_id=job_id)

        context = {
            "job_id": job.job_id,
            "job_type": job.get_job_type_display(),
            "status": job.status,
            "progress": job.progress_percentage,
            "error_message": job.error_message,
        }

        if job.status == "success":
            context["result"] = job.result_data

        response = render(
            request,
            template_name="dashboard/partials/job-progress.html",
            context=context,
        )

        # Fire distributionReady event so the table refreshes automatically
        if job.status == "success" and job.job_type == "distribution":
            response["HX-Trigger"] = "distributionReady"

        return response
    except BackgroundJob.DoesNotExist:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "Job not found"},
        )


@login_required(login_url="login")
@admin_required
def distribution_statistics(request: HttpRequest) -> HttpResponse:
    """
    Display comprehensive distribution statistics and optimization results.
    """
    date = request.GET.get("date")
    period = request.GET.get("period")

    if not date or not period:
        messages.warning(request, "Please select a date and period to view statistics.")
        return redirect("distribution")

    # Get distribution statistics
    from .utils import get_distribution_statistics

    try:
        stats = get_distribution_statistics(date, period)

        if not stats:
            messages.info(request, f"No distribution found for {date} period {period}.")
            return redirect("distribution")

        # Calculate additional metrics
        efficiency_score = (
            (stats["total_students"] / stats["total_capacity"]) * 100
            if stats["total_capacity"] > 0
            else 0
        )
        avg_utilization = (
            sum(hall["utilization"] for hall in stats["hall_details"])
            / len(stats["hall_details"])
            if stats["hall_details"]
            else 0
        )

        # Categorize halls by utilization
        high_utilization = [h for h in stats["hall_details"] if h["utilization"] >= 80]
        medium_utilization = [
            h for h in stats["hall_details"] if 50 <= h["utilization"] < 80
        ]
        low_utilization = [h for h in stats["hall_details"] if h["utilization"] < 50]

        context = {
            "date": date,
            "period": period,
            "stats": stats,
            "efficiency_score": round(efficiency_score, 1),
            "avg_utilization": round(avg_utilization, 1),
            "high_utilization_halls": high_utilization,
            "medium_utilization_halls": medium_utilization,
            "low_utilization_halls": low_utilization,
            "optimization_grade": (
                "Excellent"
                if efficiency_score >= 85
                else "Good"
                if efficiency_score >= 70
                else "Fair"
                if efficiency_score >= 50
                else "Poor"
            ),
        }

        return render(request, "distribution_statistics.html", context)

    except Exception as e:
        messages.error(request, f"Error retrieving statistics: {str(e)}")
        return redirect("distribution")


# ----------------------------
# Upload Views
# ----------------------------


@require_POST
@login_required(login_url="login")
def upload_courses(request):
    from django.db import transaction

    settings = SystemSettings.objects.first()
    if settings.has_timetable:
        return HttpResponse(
            '<div class="alert alert-danger">Courses upload not allowed again!</div>'
        )

    data = request.FILES.get("file")

    try:
        courses_df = pd.read_csv(data)

        # Validate required columns
        required_columns = ["COURSE CODE", "COURSE TITLE", "EXAM TYPE"]
        missing_columns = [
            col for col in required_columns if col not in courses_df.columns
        ]
        if missing_columns:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={
                    "message": f"Missing required columns: {', '.join(missing_columns)}"
                },
            )

        # Check for duplicate course codes in the CSV
        duplicate_codes = courses_df[
            courses_df.duplicated(subset=["COURSE CODE"], keep=False)
        ]["COURSE CODE"].tolist()
        if duplicate_codes:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={
                    "message": f"Duplicate course codes found in file: {', '.join(set(duplicate_codes))}"
                },
            )

        courses = courses_df.to_dict()
        created_count = 0
        updated_count = 0

        with transaction.atomic():
            for key in courses["COURSE CODE"]:
                course, created = Course.objects.get_or_create(
                    code=courses["COURSE CODE"][key],
                    defaults={
                        "name": courses["COURSE TITLE"][key],
                        "exam_type": courses["EXAM TYPE"][key],
                    },
                )
                if created:
                    created_count += 1
                else:
                    # Update existing course
                    course.name = courses["COURSE TITLE"][key]
                    course.exam_type = courses["EXAM TYPE"][key]
                    course.save()
                    updated_count += 1

        message = f"Courses processed successfully! Created: {created_count}, Updated: {updated_count}"
        return render(
            request,
            template_name="dashboard/partials/alert-success.html",
            context={"message": message},
        )
    except Exception as e:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": f"Upload error: {str(e)}. Please check your file format and try again."
            },
        )


@require_POST
@login_required(login_url="login")
def upload_classes(request, dept_slug):
    from django.db import transaction

    settings = SystemSettings.objects.first()
    if settings.has_timetable:
        return HttpResponse(
            '<div class="alert alert-danger">Classes upload not allowed again!</div>'
        )

    department = get_object_or_404(Department, slug=dept_slug)
    data = request.FILES.get("file")

    try:
        dept_df = pd.read_csv(data)

        # Validate required columns
        required_columns = ["Name", "Size"]
        missing_columns = [
            col for col in required_columns if col not in dept_df.columns
        ]
        if missing_columns:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={
                    "message": f"Missing required columns: {', '.join(missing_columns)}"
                },
            )

        # Check for duplicate class names in the CSV
        duplicate_names = dept_df[dept_df.duplicated(subset=["Name"], keep=False)][
            "Name"
        ].tolist()
        if duplicate_names:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={
                    "message": f"Duplicate class names found in file: {', '.join(set(duplicate_names))}"
                },
            )

        dept = dept_df.to_dict()
        created_count = 0
        updated_count = 0

        with transaction.atomic():
            for key in dept["Name"]:
                cls, created = Class.objects.get_or_create(
                    name=dept["Name"][key],
                    department=department,
                    defaults={"size": dept["Size"][key]},
                )
                if created:
                    created_count += 1
                else:
                    # Update existing class size
                    cls.size = dept["Size"][key]
                    cls.save()
                    updated_count += 1

        return render(
            request,
            template_name="dashboard/partials/alert-success.html",
            context={
                "message": f"Classes processed successfully! Created: {created_count}, Updated: {updated_count}"
            },
        )
    except Exception as e:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": f"Upload error: {str(e)}. Please check your file format and try again."
            },
        )


@require_POST
@login_required(login_url="login")
@admin_required
def upload_departments(request):
    from django.db import transaction

    settings = SystemSettings.objects.first()
    if settings.has_timetable:
        return HttpResponse(
            '<div class="alert alert-danger">Departments upload not allowed again!</div>'
        )

    data = request.FILES.get("file")

    try:
        dept_df = pd.read_csv(data)

        # Validate required columns
        required_columns = ["Name", "Code"]
        missing_columns = [
            col for col in required_columns if col not in dept_df.columns
        ]
        if missing_columns:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={
                    "message": f"Missing required columns: {', '.join(missing_columns)}"
                },
            )

        # Check for duplicate department codes in the CSV
        duplicate_codes = dept_df[dept_df.duplicated(subset=["Code"], keep=False)][
            "Code"
        ].tolist()
        if duplicate_codes:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={
                    "message": f"Duplicate department codes found in file: {', '.join(set(duplicate_codes))}"
                },
            )

        dept = dept_df.to_dict()
        created_count = 0
        updated_count = 0

        with transaction.atomic():
            for key in dept["Code"]:
                department, created = Department.objects.get_or_create(
                    slug=dept["Code"][key],
                    defaults={"name": dept["Name"][key]},
                )
                if created:
                    created_count += 1
                else:
                    # Update existing department name
                    department.name = dept["Name"][key]
                    department.save()
                    updated_count += 1

        return render(
            request,
            template_name="dashboard/partials/alert-success.html",
            context={
                "message": f"Departments processed successfully! Created: {created_count}, Updated: {updated_count}"
            },
        )
    except Exception as e:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": f"Upload error: {str(e)}. Please check your file format and try again."
            },
        )


@require_POST
@login_required(login_url="login")
def upload_class_courses(request, id):
    settings = SystemSettings.objects.first()
    if settings.has_timetable:
        return HttpResponse(
            '<div class="alert alert-danger">Class courses upload not allowed again!</div>'
        )
    cls = get_object_or_404(Class, id=id)

    # Validation 2: Check if any courses exist in the system
    if not Course.objects.exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": "No courses found in the system. Admin must upload the institutional course catalog before class courses can be uploaded."
            },
        )

    # Get uploaded file
    data = request.FILES.get("file")
    df = pd.read_csv(data)
    courses_data = df.to_dict()
    # Get all existing course codes from the system
    existing_course_codes = set(Course.objects.values_list("code", flat=True))
    # Check each course code in the CSV
    invalid_codes = []
    for key in courses_data["COURSE CODE"]:
        course_code = courses_data["COURSE CODE"][key]
        if course_code not in existing_course_codes:
            invalid_codes.append(course_code)
    # If any invalid codes found, return error
    if invalid_codes:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": f"The following course codes do not exist in the system: {', '.join(invalid_codes)}"
            },
        )
    # All codes are valid, proceed with upload
    for key in courses_data["COURSE CODE"]:
        course_code = courses_data["COURSE CODE"][key]
        course = Course.objects.get(code=course_code)
        cls.courses.add(course)
    cls.save()
    return render(
        request,
        template_name="dashboard/partials/alert-success.html",
        context={"message": "Courses uploaded successfully"},
    )


@require_POST
@login_required(login_url="login")
def upload_class_students(request, id):
    from django.db import transaction

    print(f"Uploading students for class {id}")

    settings = SystemSettings.objects.first()
    if settings.has_timetable:
        return HttpResponse(
            '<div class="alert alert-danger">Class students upload not allowed again!</div>'
        )

    cls = get_object_or_404(Class, id=id)
    data = request.FILES.get("file")

    try:
        # Read CSV efficiently with optimized settings
        students_df = pd.read_csv(
            data,
            dtype={
                "MATRIC NUMBER": "string",
                "FIRSTNAME": "string",
                "LASTNAME": "string",
                "EMAIL": "string",
                "PHONE NUMBER": "string",
            },
            engine="c",  # Use C engine for faster parsing
        )
        print(f"Uploading students for class {id}")

        # Remove any whitespace and convert to string
        students_df = students_df.astype(str).apply(lambda x: x.str.strip())

        # Validation: Check if number of students matches class size
        total_students_in_file = len(students_df)
        print(f"Uploading students for class {id}")

        # if total_students_in_file != cls.size:
        #     return render(
        #         request,
        #         template_name="dashboard/partials/alert-error.html",
        #         context={
        #             "message": f"Student count mismatch. Class size is {cls.size} but you are uploading {total_students_in_file} students. Please ensure the number of students matches the class size."
        #         },
        #     )

        # Validate required columns
        required_columns = [
            "MATRIC NUMBER",
            "FIRSTNAME",
            "LASTNAME",
            "EMAIL",
            "PHONE NUMBER",
        ]
        missing_columns = [
            col for col in required_columns if col not in students_df.columns
        ]
        if missing_columns:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={
                    "message": f"Missing required columns: {', '.join(missing_columns)}"
                },
            )
        print(f"Uploading {total_students_in_file} students for class {cls.name}")
        # Check for duplicate matric numbers in the file
        duplicate_matrics = students_df[
            students_df.duplicated(subset=["MATRIC NUMBER"], keep=False)
        ]["MATRIC NUMBER"].tolist()
        if duplicate_matrics:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={
                    "message": f"Duplicate matric numbers found in file: {', '.join(duplicate_matrics[:5])}{'...' if len(duplicate_matrics) > 5 else ''}"
                },
            )

        # Convert DataFrame to records for faster processing
        student_records = students_df.to_dict("records")
        matric_numbers = [record["MATRIC NUMBER"] for record in student_records]

        # Get existing students in one optimized query
        print(f"Checking for {len(matric_numbers)} existing students")

        existing_students = {}
        for student in Student.objects.filter(matric_no__in=matric_numbers).only(
            "id",
            "matric_no",
            "first_name",
            "last_name",
            "email",
            "phone",
            "department_id",
            "level_id",
        ):
            existing_students[student.matric_no] = student

        # Process in chunks for better memory management
        chunk_size = 250
        total_created = 0
        total_updated = 0

        with transaction.atomic():
            for i in range(0, len(student_records), chunk_size):
                chunk = student_records[i : i + chunk_size]
                students_to_create = []
                students_to_update = []

                for record in chunk:
                    matric_no = record["MATRIC NUMBER"]

                    if matric_no in existing_students:
                        # Prepare for bulk update
                        student = existing_students[matric_no]
                        student.first_name = record["FIRSTNAME"]
                        student.last_name = record["LASTNAME"]
                        student.email = record["EMAIL"]
                        student.phone = record["PHONE NUMBER"]
                        student.department = cls.department
                        student.level = cls
                        students_to_update.append(student)
                    else:
                        # Prepare for bulk create
                        students_to_create.append(
                            Student(
                                matric_no=matric_no,
                                first_name=record["FIRSTNAME"],
                                last_name=record["LASTNAME"],
                                email=record["EMAIL"],
                                phone=record["PHONE NUMBER"],
                                department=cls.department,
                                level=cls,
                            )
                        )

                # Bulk create new students for this chunk
                if students_to_create:
                    Student.objects.bulk_create(students_to_create, batch_size=250)
                    total_created += len(students_to_create)

                # Bulk update existing students for this chunk
                if students_to_update:
                    Student.objects.bulk_update(
                        students_to_update,
                        [
                            "first_name",
                            "last_name",
                            "email",
                            "phone",
                            "department",
                            "level",
                        ],
                        batch_size=250,
                    )
                    total_updated += len(students_to_update)

            created_count = total_created
            updated_count = total_updated
            print(f"Processed chunk: Created {created_count}, Updated {updated_count}")

            if created_count > 0 or updated_count > 0:
                message = f"Students processed successfully! Created: {created_count}, Updated: {updated_count}"
                return render(
                    request,
                    template_name="dashboard/partials/alert-success.html",
                    context={"message": message},
                )
            else:
                return render(
                    request,
                    template_name="dashboard/partials/alert-error.html",
                    context={
                        "message": "No students were processed. Please check your data."
                    },
                )

    except Exception as e:
        print(f"Error in upload_students: {str(e)}")

        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": f"Upload error: {str(e)}. Please check your file format and try again."
            },
        )


@require_POST
@login_required(login_url="login")
@admin_required
def upload_halls(request):
    from django.db import transaction

    settings = SystemSettings.objects.first()
    if settings.has_timetable:
        return HttpResponse(
            '<div class="alert alert-danger">Halls upload not allowed again!</div>'
        )

    data = request.FILES.get("file")

    try:
        halls_df = pd.read_csv(data)

        # Validate required columns
        required_columns = [
            "EXAM VENUE",
            "CAPACITY",
            "MAX STUDENTS",
            "MIN COURSES",
            "ROWS",
            "COLS",
        ]
        missing_columns = [
            col for col in required_columns if col not in halls_df.columns
        ]
        if missing_columns:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={
                    "message": f"Missing required columns: {', '.join(missing_columns)}"
                },
            )

        # Check for duplicate hall names in the CSV
        duplicate_halls = halls_df[
            halls_df.duplicated(subset=["EXAM VENUE"], keep=False)
        ]["EXAM VENUE"].tolist()
        if duplicate_halls:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={
                    "message": f"Duplicate hall names found in file: {', '.join(set(duplicate_halls))}"
                },
            )

        halls = halls_df.to_dict()
        created_count = 0
        updated_count = 0

        with transaction.atomic():
            for key in halls["EXAM VENUE"]:
                hall, created = Hall.objects.get_or_create(
                    name=halls["EXAM VENUE"][key],
                    defaults={
                        "capacity": halls["CAPACITY"][key],
                        "max_students": halls["MAX STUDENTS"][key],
                        "min_courses": halls["MIN COURSES"][key],
                        "rows": halls["ROWS"][key],
                        "columns": halls["COLS"][key],
                    },
                )
                if created:
                    created_count += 1
                else:
                    # Update existing hall
                    hall.capacity = halls["CAPACITY"][key]
                    hall.max_students = halls["MAX STUDENTS"][key]
                    hall.min_courses = halls["MIN COURSES"][key]
                    hall.rows = halls["ROWS"][key]
                    hall.columns = halls["COLS"][key]
                    hall.save()
                    updated_count += 1

        message = f"Halls processed successfully! Created: {created_count}, Updated: {updated_count}"
        return render(
            request,
            template_name="dashboard/partials/alert-success.html",
            context={"message": message},
        )
    except Exception as e:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": f"Upload error: {str(e)}. Please check your file format and try again."
            },
        )


@require_POST
@login_required(login_url="login")
@admin_required
def enable_bulk_upload(request):
    """Enable bulk upload by setting has_timetable to False"""
    settings = SystemSettings.objects.first()
    if settings:
        settings.has_timetable = False
        settings.save()
        messages.success(request, "Bulk upload has been enabled successfully!")
    else:
        messages.error(request, "System settings not found!")
    return redirect("bulk-upload")


def bulk_upload(request):
    settings = SystemSettings.objects.first()
    if settings.has_timetable:
        messages.warning(
            request,
            "Bulk upload is currently disabled because a timetable has been generated. "
            '<br><small class="text-muted">To enable bulk upload again, you can either:</small><br>'
            '<a href="#" onclick="enableBulkUpload()" class="btn btn-sm btn-primary mt-2">Enable Bulk Upload</a> '
            '<a href="/reset/" class="btn btn-sm btn-danger mt-2" onclick="return confirm(\'This will reset the entire system. Are you sure?\')">Reset System</a>'
            "<br><br><script>function enableBulkUpload(){fetch('/enable-bulk-upload/', {method: 'POST', headers: {'X-CSRFToken': document.querySelector('[name=csrfmiddlewaretoken]').value}}).then(response => response.text()).then(data => {location.reload();});}</script>",
        )
        return render(request, template_name="dashboard/upload.html")

    if request.method == "POST":
        try:
            file = request.FILES["file"]
            upload_type = request.POST["upload_type"]
            result = handle_uploaded_file(file, upload_type)

            # Dynamic success messages based on upload type
            success_messages = {
                "courses": "Course files uploaded and processed successfully!",
                "classes": "Class files uploaded and processed successfully!",
                "students": "Student files uploaded and processed successfully!",
                "halls": "Hall files uploaded and processed successfully!",
            }

            message = success_messages.get(upload_type, "Files uploaded successfully!")

            # If result contains detailed information (like from student processing)
            if isinstance(result, dict) and "message" in result:
                message = result["message"]

            messages.success(request, message)
            return redirect("bulk-upload")

        except ValueError as e:
            # Handle validation errors with detailed formatting
            error_message = str(e)

            # Check if it's a multi-line error message
            if "\n" in error_message:
                # Split into main message and details
                lines = error_message.split("\n")
                main_message = lines[0]
                details = "\n".join(lines[1:]) if len(lines) > 1 else ""

                formatted_message = f"<strong>{main_message}</strong>"
                if details.strip():
                    # Convert newlines to HTML breaks and preserve formatting
                    formatted_details = details.replace("\n", "<br>").replace(
                        "  •", "&nbsp;&nbsp;•"
                    )
                    formatted_message += f'<br><br><div style="font-family: monospace; font-size: 0.9em; background: #f8f9fa; padding: 10px; border-radius: 4px; margin-top: 8px;">{formatted_details}</div>'

                messages.error(request, formatted_message)
            else:
                messages.error(
                    request, f"<strong>Validation Error:</strong><br>{error_message}"
                )
            return render(request, template_name="dashboard/upload.html")

        except Exception as e:
            # Handle unexpected errors
            messages.error(
                request,
                f'<strong>Upload Error:</strong><br>An unexpected error occurred: {str(e)}<br><small class="text-muted">Please check your file format and try again.</small>',
            )
            return render(request, template_name="dashboard/upload.html")
    else:
        return render(request, template_name="dashboard/upload.html")


@login_required(login_url="login")
@admin_required
def job_monitor_view(request):
    """
    View to monitor all background jobs
    """
    # Get filter parameters
    status_filter = request.GET.get("status", "")
    job_type_filter = request.GET.get("job_type", "")

    # Base query - if staff, show all jobs; otherwise, show only user's jobs
    if request.user.is_staff:
        jobs = BackgroundJob.objects.all()
    else:
        jobs = BackgroundJob.objects.filter(created_by=request.user)

    # Apply filters
    if status_filter:
        jobs = jobs.filter(status=status_filter)
    if job_type_filter:
        jobs = jobs.filter(job_type=job_type_filter)

    # Order by most recent first
    jobs = jobs.select_related("created_by").order_by("-started_at")

    # Paginate
    paginator = Paginator(jobs, 20)
    page_number = request.GET.get("page", 1)
    page_obj = paginator.get_page(page_number)

    # Count running jobs for the current user
    running_jobs_count = BackgroundJob.objects.filter(
        created_by=request.user, status__in=["pending", "running"]
    ).count()

    context = {
        "page_obj": page_obj,
        "status_filter": status_filter,
        "job_type_filter": job_type_filter,
        "running_jobs_count": running_jobs_count,
    }

    if request.htmx:
        template_name = "dashboard/pages/jobs.html"
    else:
        template_name = "dashboard/jobs.html"

    return render(request, template_name, context)


@login_required(login_url="login")
@admin_required
def job_detail_view(request, job_id):
    """
    View detailed information about a specific background job
    """
    # Get the job - if staff, can view any job; otherwise, only own jobs
    if request.user.is_staff:
        job = get_object_or_404(BackgroundJob, job_id=job_id)
    else:
        job = get_object_or_404(BackgroundJob, job_id=job_id, created_by=request.user)

    context = {
        "job": job,
    }

    if request.htmx:
        template_name = "dashboard/pages/job-details.html"
    else:
        template_name = "dashboard/job-detail.html"

    return render(request, template_name, context)


@login_required(login_url="login")
@admin_required
@require_POST
def job_delete_view(request, job_id):
    """
    Delete a background job record
    """
    # Get the job - if staff, can delete any job; otherwise, only own jobs
    if request.user.is_staff:
        job = get_object_or_404(BackgroundJob, job_id=job_id)
    else:
        job = get_object_or_404(BackgroundJob, job_id=job_id, created_by=request.user)

    # Only allow deletion of completed or failed jobs
    if job.status in ["success", "failed"]:
        job.delete()
        messages.success(request, "Job deleted successfully.")
    else:
        messages.error(request, "Cannot delete a running job.")

    return redirect("job_monitor")


@login_required(login_url="login")
@admin_required
@require_POST
def job_retry_view(request, job_id):
    """
    Retry a failed background job
    """
    # Get the job - if staff, can retry any job; otherwise, only own jobs
    if request.user.is_staff:
        job = get_object_or_404(BackgroundJob, job_id=job_id)
    else:
        job = get_object_or_404(BackgroundJob, job_id=job_id, created_by=request.user)

    # Only allow retry of failed jobs
    if job.status != "failed":
        messages.error(request, "Can only retry failed jobs.")
        return redirect("job_monitor")

    # Import tasks
    # Create a new job
    import uuid

    from django.db import transaction

    from .tasks import (
        generate_allocation_task,
        generate_distribution_task,
        generate_timetable_task,
    )

    new_job_id = str(uuid.uuid4())
    new_job = BackgroundJob.objects.create(
        job_id=new_job_id,
        job_type=job.job_type,
        status="pending",
        created_by=request.user,
        params=job.params,
    )

    # Trigger the appropriate task based on job type (after commit)
    def trigger_retry_task():
        if job.job_type == "timetable":
            params = job.params
            generate_timetable_task.apply_async(
                args=[
                    new_job_id,
                    request.user.id,
                    params.get("start_date"),
                    params.get("end_date"),
                ]
            )
        elif job.job_type == "distribution":
            params = job.params
            generate_distribution_task.apply_async(
                args=[
                    new_job_id,
                    request.user.id,
                    params.get("date"),
                    params.get("period"),
                ]
            )
        elif job.job_type == "allocation":
            params = job.params
            generate_allocation_task.apply_async(
                args=[
                    new_job_id,
                    request.user.id,
                    params.get("date"),
                    params.get("period"),
                ]
            )

    transaction.on_commit(trigger_retry_task)

    messages.success(request, f"Job retried successfully. New job ID: {new_job_id}")
    return redirect("job_detail", job_id=new_job_id)


# --------------------------
#  Department CRUD Views
# --------------------------


@require_POST
@login_required(login_url="login")
@admin_required
def create_department(request):
    """Create a new department"""
    name = request.POST.get("name", "").strip()
    slug = request.POST.get("slug", "").strip().upper()

    if not name or not slug:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "Department name and code are required"},
        )

    if Department.objects.filter(slug=slug).exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": f"Department with code '{slug}' already exists"},
        )

    if Department.objects.filter(name__iexact=name).exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": f"Department with name '{name}' already exists"},
        )

    Department.objects.create(name=name, slug=slug)
    return render(
        request,
        template_name="dashboard/partials/alert-success.html",
        context={"message": "Department created successfully"},
    )


@login_required(login_url="login")
@admin_required
def edit_department(request, slug):
    """Edit a department"""
    department = get_object_or_404(Department, slug=slug)

    if request.method == "POST":
        name = request.POST.get("name", "").strip()
        new_slug = request.POST.get("slug", "").strip().upper()

        if not name or not new_slug:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={"message": "Department name and code are required"},
            )

        # Check if new slug conflicts with another department
        if (
            new_slug != department.slug
            and Department.objects.filter(slug=new_slug).exists()
        ):
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={
                    "message": f"Department with code '{new_slug}' already exists"
                },
            )

        # Check if new name conflicts with another department
        if (
            name.lower() != department.name.lower()
            and Department.objects.filter(name__iexact=name)
            .exclude(id=department.id)
            .exists()
        ):
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={"message": f"Department with name '{name}' already exists"},
            )

        department.name = name
        department.slug = new_slug
        department.save()

        return render(
            request,
            template_name="dashboard/partials/alert-success.html",
            context={"message": "Department updated successfully"},
        )

    # GET request - return modal content
    context = {"department": department}
    return render(request, "dashboard/partials/edit-department-modal.html", context)


@require_POST
@login_required(login_url="login")
@admin_required
def delete_department(request, slug):
    """Delete a department"""
    department = get_object_or_404(Department, slug=slug)

    # Check if department has associated classes
    if Class.objects.filter(department=department).exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": "Cannot delete department with associated classes. Delete the classes first."
            },
        )

    # Check if department has associated users
    if User.objects.filter(department=department).exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": "Cannot delete department with associated users. Reassign the users first."
            },
        )

    department.delete()
    return render(
        request,
        template_name="dashboard/partials/alert-success.html",
        context={"message": "Department deleted successfully"},
    )


# --------------------------
#  User Management CRUD Views
# --------------------------


@login_required(login_url="login")
@admin_required
def edit_user(request, user_id):
    """Edit a user"""
    user = get_object_or_404(User, id=user_id)
    departments = Department.objects.all()

    if request.method == "POST":
        first_name = request.POST.get("first_name", "").strip()
        last_name = request.POST.get("last_name", "").strip()
        email = request.POST.get("email", "").strip()
        department_slug = request.POST.get("department", "")
        is_staff = request.POST.get("is_staff") == "on"

        if not first_name or not last_name or not email:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={"message": "First name, last name and email are required"},
            )

        # Check if email conflicts with another user
        if email != user.email and User.objects.filter(email=email).exists():
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={"message": "Email already exists"},
            )

        user.first_name = first_name
        user.last_name = last_name
        user.email = email
        user.is_staff = is_staff

        if department_slug and not is_staff:
            department = Department.objects.filter(slug=department_slug).first()
            user.department = department
        elif is_staff:
            user.department = None

        user.save()

        return render(
            request,
            template_name="dashboard/partials/alert-success.html",
            context={"message": "User updated successfully"},
        )

    # GET request - return modal content
    context = {"user_obj": user, "departments": departments}
    return render(request, "dashboard/partials/edit-user-modal.html", context)


@require_POST
@login_required(login_url="login")
@admin_required
def delete_user(request, user_id):
    """Delete a user"""
    user = get_object_or_404(User, id=user_id)

    # Prevent deleting yourself
    if user.id == request.user.id:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "You cannot delete your own account"},
        )

    # Prevent deleting the last admin
    if user.is_staff and User.objects.filter(is_staff=True).count() <= 1:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "Cannot delete the last admin user"},
        )

    user.delete()
    return render(
        request,
        template_name="dashboard/partials/alert-success.html",
        context={"message": "User deleted successfully"},
    )


@login_required(login_url="login")
@admin_required
def change_user_password(request, user_id):
    """Change a user's password"""
    user = get_object_or_404(User, id=user_id)

    if request.method == "POST":
        new_password = request.POST.get("new_password", "")
        confirm_password = request.POST.get("confirm_password", "")

        if not new_password:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={"message": "Password is required"},
            )

        if new_password != confirm_password:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={"message": "Passwords do not match"},
            )

        if len(new_password) < 8:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={"message": "Password must be at least 8 characters"},
            )

        user.set_password(new_password)
        user.save()

        return render(
            request,
            template_name="dashboard/partials/alert-success.html",
            context={"message": "Password changed successfully"},
        )

    # GET request - return modal content
    context = {"user_obj": user}
    return render(request, "dashboard/partials/change-password-modal.html", context)


# --------------------------
#  Class CRUD Views
# --------------------------


@require_POST
@login_required(login_url="login")
@admin_required
def create_class(request, dept_slug):
    """Create a new class in a department"""
    department = get_object_or_404(Department, slug=dept_slug)

    name = request.POST.get("name", "").strip()
    size = request.POST.get("size", "0").strip()

    if not name:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "Class name is required"},
        )

    try:
        size = int(size)
        if size < 0:
            raise ValueError("Size must be positive")
    except ValueError:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "Class size must be a valid positive number"},
        )

    # Check if class already exists in this department
    if Class.objects.filter(name__iexact=name, department=department).exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": f"Class '{name}' already exists in {department.name}"},
        )

    Class.objects.create(name=name, size=size, department=department)
    return render(
        request,
        template_name="dashboard/partials/alert-success.html",
        context={"message": "Class created successfully"},
    )


@login_required(login_url="login")
@admin_required
def edit_class(request, class_id):
    """Edit a class"""
    cls = get_object_or_404(Class, id=class_id)

    if request.method == "POST":
        name = request.POST.get("name", "").strip()
        size = request.POST.get("size", "0").strip()

        if not name:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={"message": "Class name is required"},
            )

        try:
            size = int(size)
            if size < 0:
                raise ValueError("Size must be positive")
        except ValueError:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={"message": "Class size must be a valid positive number"},
            )

        # Check if class name conflicts with another class in the same department
        if (
            name.lower() != cls.name.lower()
            and Class.objects.filter(name__iexact=name, department=cls.department)
            .exclude(id=cls.id)
            .exists()
        ):
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={
                    "message": f"Class '{name}' already exists in {cls.department.name}"
                },
            )

        cls.name = name
        cls.size = size
        cls.save()

        return render(
            request,
            template_name="dashboard/partials/alert-success.html",
            context={"message": "Class updated successfully"},
        )

    # GET request - return modal content
    context = {"class_obj": cls}
    return render(request, "dashboard/partials/edit-class-modal.html", context)


@require_POST
@login_required(login_url="login")
@admin_required
def delete_class(request, class_id):
    """Delete a class"""
    cls = get_object_or_404(Class, id=class_id)

    # Check if class has students
    if Student.objects.filter(level=cls).exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": "Cannot delete class with students. Delete the students first."
            },
        )

    # Check if class has timetable entries
    if TimeTable.objects.filter(class_obj=cls).exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": "Cannot delete class with timetable entries. Clear the timetable first."
            },
        )

    cls.delete()
    return render(
        request,
        template_name="dashboard/partials/alert-success.html",
        context={"message": "Class deleted successfully"},
    )


# --------------------------
#  Course CRUD Views (within Class context)
# --------------------------


@require_POST
@login_required(login_url="login")
@admin_required
def add_course_to_class(request, class_id):
    """Add a course to a class (create if doesn't exist, or add existing)"""
    cls = get_object_or_404(Class, id=class_id)

    name = request.POST.get("name", "").strip()
    code = request.POST.get("code", "").strip().upper()
    exam_type = request.POST.get("exam_type", "PBE").strip()

    if not name or not code:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "Course name and code are required"},
        )

    if exam_type not in ["PBE", "CBE"]:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "Invalid exam type. Must be PBE or CBE"},
        )

    # Check if course already exists
    course = Course.objects.filter(code__iexact=code).first()

    if course:
        # Course exists, check if already in class
        if course in cls.courses.all():
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={
                    "message": f"Course '{code}' is already assigned to this class"
                },
            )
        # Add existing course to class
        cls.courses.add(course)
    else:
        # Create new course and add to class
        course = Course.objects.create(name=name, code=code, exam_type=exam_type)
        cls.courses.add(course)

    return render(
        request,
        template_name="dashboard/partials/alert-success.html",
        context={"message": f"Course '{code}' added to class successfully"},
    )


@login_required(login_url="login")
@admin_required
def edit_class_course(request, class_id, course_id):
    """Edit a course within class context"""
    cls = get_object_or_404(Class, id=class_id)
    course = get_object_or_404(Course, id=course_id)

    # Verify course is in this class
    if course not in cls.courses.all():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "Course not found in this class"},
        )

    if request.method == "POST":
        name = request.POST.get("name", "").strip()
        code = request.POST.get("code", "").strip().upper()
        exam_type = request.POST.get("exam_type", "PBE").strip()

        if not name or not code:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={"message": "Course name and code are required"},
            )

        if exam_type not in ["PBE", "CBE"]:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={"message": "Invalid exam type. Must be PBE or CBE"},
            )

        # Check if course code conflicts with another course
        if (
            code.lower() != course.code.lower()
            and Course.objects.filter(code__iexact=code).exclude(id=course.id).exists()
        ):
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={"message": f"Course with code '{code}' already exists"},
            )

        course.name = name
        course.code = code
        course.exam_type = exam_type
        course.save()

        return render(
            request,
            template_name="dashboard/partials/alert-success.html",
            context={"message": "Course updated successfully"},
        )

    # GET request - return modal content
    context = {"course": course, "class": cls}
    return render(request, "dashboard/partials/edit-course-modal.html", context)


@require_POST
@login_required(login_url="login")
@admin_required
def remove_course_from_class(request, class_id, course_id):
    """Remove a course from a class (does not delete the course itself)"""
    cls = get_object_or_404(Class, id=class_id)
    course = get_object_or_404(Course, id=course_id)

    # Verify course is in this class
    if course not in cls.courses.all():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "Course not found in this class"},
        )

    # Check if course is used in timetable
    if TimeTable.objects.filter(course=course).exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": "Cannot remove course with timetable entries. Clear the timetable first."
            },
        )

    # Remove course from class (not deleting the course)
    cls.courses.remove(course)

    return render(
        request,
        template_name="dashboard/partials/alert-success.html",
        context={"message": f"Course '{course.code}' removed from class successfully"},
    )


# --------------------------
#  Global Course CRUD Views (for Courses page in sidebar)
# --------------------------


@require_POST
@login_required(login_url="login")
@admin_required
def create_course(request):
    """Create a new course (global)"""
    name = request.POST.get("name", "").strip()
    code = request.POST.get("code", "").strip().upper()
    exam_type = request.POST.get("exam_type", "PBE").strip()

    if not name or not code:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "Course name and code are required"},
        )

    if exam_type not in ["PBE", "CBE"]:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "Invalid exam type. Must be PBE or CBE"},
        )

    if Course.objects.filter(code__iexact=code).exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": f"Course with code '{code}' already exists"},
        )

    Course.objects.create(name=name, code=code, exam_type=exam_type)
    return render(
        request,
        template_name="dashboard/partials/alert-success.html",
        context={"message": "Course created successfully"},
    )


@login_required(login_url="login")
@admin_required
def edit_course(request, course_id):
    """Edit a course (global)"""
    course = get_object_or_404(Course, id=course_id)

    if request.method == "POST":
        name = request.POST.get("name", "").strip()
        code = request.POST.get("code", "").strip().upper()
        exam_type = request.POST.get("exam_type", "PBE").strip()

        if not name or not code:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={"message": "Course name and code are required"},
            )

        if exam_type not in ["PBE", "CBE"]:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={"message": "Invalid exam type. Must be PBE or CBE"},
            )

        # Check if course code conflicts with another course
        if (
            code.lower() != course.code.lower()
            and Course.objects.filter(code__iexact=code).exclude(id=course.id).exists()
        ):
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={"message": f"Course with code '{code}' already exists"},
            )

        course.name = name
        course.code = code
        course.exam_type = exam_type
        course.save()

        return render(
            request,
            template_name="dashboard/partials/alert-success.html",
            context={"message": "Course updated successfully"},
        )

    # GET request - return modal content
    context = {"course": course}
    return render(request, "dashboard/partials/edit-course-global-modal.html", context)


@require_POST
@login_required(login_url="login")
@admin_required
def delete_course(request, course_id):
    """Delete a course (global)"""
    course = get_object_or_404(Course, id=course_id)

    # Check if course is used in timetable
    if TimeTable.objects.filter(course=course).exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": "Cannot delete course with timetable entries. Clear the timetable first."
            },
        )

    # Check if course is assigned to any class
    if Class.objects.filter(courses=course).exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": "Cannot delete course assigned to classes. Remove from classes first."
            },
        )

    course.delete()
    return render(
        request,
        template_name="dashboard/partials/alert-success.html",
        context={"message": "Course deleted successfully"},
    )


# --------------------------
#  Student CRUD Views (within Class context)
# --------------------------


@require_POST
@login_required(login_url="login")
@admin_required
def create_student(request, class_id):
    """Create a new student in a class"""
    cls = get_object_or_404(Class, id=class_id)

    first_name = request.POST.get("first_name", "").strip()
    last_name = request.POST.get("last_name", "").strip()
    matric_no = request.POST.get("matric_no", "").strip().upper()
    email = request.POST.get("email", "").strip()
    phone = request.POST.get("phone", "").strip()

    if not first_name or not last_name or not matric_no:
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={"message": "First name, last name and matric number are required"},
        )

    if Student.objects.filter(matric_no=matric_no).exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": f"Student with matric number '{matric_no}' already exists"
            },
        )

    Student.objects.create(
        first_name=first_name,
        last_name=last_name,
        matric_no=matric_no,
        email=email,
        phone=phone,
        level=cls,
        department=cls.department,
    )

    # Update class size
    cls.size = Student.objects.filter(level=cls, department=cls.department).count()
    cls.save()

    return render(
        request,
        template_name="dashboard/partials/alert-success.html",
        context={"message": "Student created successfully"},
    )


@login_required(login_url="login")
@admin_required
def edit_student(request, student_id):
    """Edit a student"""
    student = get_object_or_404(Student, id=student_id)

    if request.method == "POST":
        first_name = request.POST.get("first_name", "").strip()
        last_name = request.POST.get("last_name", "").strip()
        matric_no = request.POST.get("matric_no", "").strip().upper()
        email = request.POST.get("email", "").strip()
        phone = request.POST.get("phone", "").strip()

        if not first_name or not last_name or not matric_no:
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={
                    "message": "First name, last name and matric number are required"
                },
            )

        # Check if matric number conflicts with another student
        if (
            matric_no != student.matric_no
            and Student.objects.filter(matric_no=matric_no).exists()
        ):
            return render(
                request,
                template_name="dashboard/partials/alert-error.html",
                context={
                    "message": f"Student with matric number '{matric_no}' already exists"
                },
            )

        student.first_name = first_name
        student.last_name = last_name
        student.matric_no = matric_no
        student.email = email
        student.phone = phone
        student.save()

        return render(
            request,
            template_name="dashboard/partials/alert-success.html",
            context={"message": "Student updated successfully"},
        )

    # GET request - return modal content
    context = {"student": student}
    return render(request, "dashboard/partials/edit-student-modal.html", context)


@require_POST
@login_required(login_url="login")
@admin_required
def delete_student(request, student_id):
    """Delete a student"""
    student = get_object_or_404(Student, id=student_id)
    cls = student.level

    # Check if student has seat arrangements
    if SeatArrangement.objects.filter(student=student).exists():
        return render(
            request,
            template_name="dashboard/partials/alert-error.html",
            context={
                "message": "Cannot delete student with seat arrangements. Clear allocations first."
            },
        )

    student.delete()

    # Update class size
    cls.size = Student.objects.filter(level=cls, department=cls.department).count()
    cls.save()

    return render(
        request,
        template_name="dashboard/partials/alert-success.html",
        context={"message": "Student deleted successfully"},
    )
